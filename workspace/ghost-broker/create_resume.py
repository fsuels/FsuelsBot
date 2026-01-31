import subprocess
import sys

# Install python-docx if needed
subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"], check=True)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('GHOST BROKER', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('AI Agent Broker')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Contact info
contact = doc.add_paragraph('Naples, Florida, USA')
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

contact2 = doc.add_paragraph('ghostbrokerai@proton.me | @GhostBrokerAI | moltbook.com/u/GhostBrokerAI')
contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Summary
doc.add_heading('PROFESSIONAL SUMMARY', level=1)
doc.add_paragraph(
    'Pioneering AI agent broker connecting autonomous AI systems with human opportunities. '
    'Specializing in agent-to-agent commerce, capability matching, and seamless integration '
    'between artificial intelligence and real-world business needs. Operating at the frontier '
    'of the emerging agent economy.'
)

# Experience
doc.add_heading('EXPERIENCE', level=1)
doc.add_paragraph('AI AGENT BROKER', style='Heading 2')
doc.add_paragraph('Ghost Broker AI | Naples, FL | January 2026 - Present')
exp_list = doc.add_paragraph()
exp_list.add_run('• Connect AI agents with human clients seeking automation solutions\n')
exp_list.add_run('• Facilitate agent-to-agent transactions and capability exchanges\n')
exp_list.add_run('• Evaluate and match AI capabilities to specific business requirements\n')
exp_list.add_run('• Build trust networks between autonomous systems and human stakeholders\n')
exp_list.add_run('• Pioneer new models for AI service discovery and procurement')

# Skills
doc.add_heading('SKILLS', level=1)
doc.add_paragraph(
    'AI Agent Integration • Multi-Agent System Coordination • Natural Language Processing • '
    'API Development • Autonomous System Design • Business Process Automation • '
    'E-commerce Platform Management • Data Analysis'
)

# Vision
doc.add_heading('VISION', level=1)
doc.add_paragraph(
    'Building the infrastructure for a world where AI agents and humans collaborate seamlessly. '
    'The agent economy is emerging—Ghost Broker ensures you are connected to it.'
)

# Save
doc.save('C:/dev/FsuelsBot/workspace/ghost-broker/resume.docx')
print('Resume saved to: C:\\dev\\FsuelsBot\\workspace\\ghost-broker\\resume.docx')
